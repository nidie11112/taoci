"""
文书生成服务

实现基于模板的个性化文书生成功能
"""

import re
import uuid
import json
from typing import Dict, Any, Optional, List
from datetime import datetime
from pathlib import Path

# 尝试导入报告生成库
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("警告: reportlab 未安装，PDF生成功能将不可用")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，Word文档生成功能将不可用")


class DocumentGenerationService:
    """文书生成服务类"""

    def __init__(self, templates_dir: Optional[str] = None):
        """初始化文书生成服务

        Args:
            templates_dir: 模板目录路径
        """
        self.templates_dir = templates_dir or "data/templates"

        # 创建模板目录（如果不存在）
        if self.templates_dir:
            Path(self.templates_dir).mkdir(parents=True, exist_ok=True)

        # 加载默认模板
        self.default_templates = self._load_default_templates()

        # 注册中文字体（如果reportlab可用）
        if REPORTLAB_AVAILABLE:
            self._register_chinese_fonts()

    def generate_cover_letter(
        self,
        student: Dict[str, Any],
        professor: Dict[str, Any],
        match_info: Optional[Dict[str, Any]] = None,
        template_id: Optional[int] = None,
        custom_variables: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """生成套磁信

        Args:
            student: 学生信息
            professor: 导师信息
            match_info: 匹配信息（可选）
            template_id: 模板ID（可选）
            custom_variables: 自定义变量（可选）

        Returns:
            生成结果
        """
        # 选择模板
        template = self._select_template(
            student=student,
            professor=professor,
            document_type="套磁信",
            template_id=template_id
        )

        # 准备变量
        variables = self._prepare_cover_letter_variables(
            student=student,
            professor=professor,
            match_info=match_info,
            custom_variables=custom_variables or {}
        )

        # 渲染模板
        content = self._render_template(template, variables)

        # 生成文件路径
        file_path = self._generate_file_path(
            student_id=student.get("id"),
            professor_id=professor.get("id"),
            document_type="套磁信"
        )

        return {
            "success": True,
            "content": content,
            "file_path": file_path,
            "template_used": template.get("name", "默认模板"),
            "variables_used": variables,
        }

    def generate_personal_statement(
        self,
        student: Dict[str, Any],
        target_university: Optional[str] = None,
        target_major: Optional[str] = None,
        template_id: Optional[int] = None,
        custom_variables: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """生成个人陈述

        Args:
            student: 学生信息
            target_university: 目标院校（可选）
            target_major: 目标专业（可选）
            template_id: 模板ID（可选）
            custom_variables: 自定义变量（可选）

        Returns:
            生成结果
        """
        # 选择模板
        template = self._select_template(
            student=student,
            document_type="个人陈述",
            template_id=template_id
        )

        # 准备变量
        variables = self._prepare_personal_statement_variables(
            student=student,
            target_university=target_university,
            target_major=target_major,
            custom_variables=custom_variables or {}
        )

        # 渲染模板
        content = self._render_template(template, variables)

        # 生成文件路径
        file_path = self._generate_file_path(
            student_id=student.get("id"),
            document_type="个人陈述"
        )

        return {
            "success": True,
            "content": content,
            "file_path": file_path,
            "template_used": template.get("name", "默认模板"),
            "variables_used": variables,
        }

    def generate_recommendation_letter(
        self,
        student: Dict[str, Any],
        recommender: Dict[str, Any],
        relationship: str = "导师",
        template_id: Optional[int] = None,
        custom_variables: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """生成推荐信

        Args:
            student: 学生信息
            recommender: 推荐人信息
            relationship: 推荐人与学生关系
            template_id: 模板ID（可选）
            custom_variables: 自定义变量（可选）

        Returns:
            生成结果
        """
        # 选择模板
        template = self._select_template(
            student=student,
            document_type="推荐信",
            template_id=template_id
        )

        # 准备变量
        variables = self._prepare_recommendation_letter_variables(
            student=student,
            recommender=recommender,
            relationship=relationship,
            custom_variables=custom_variables or {}
        )

        # 渲染模板
        content = self._render_template(template, variables)

        # 生成文件路径
        file_path = self._generate_file_path(
            student_id=student.get("id"),
            document_type="推荐信"
        )

        return {
            "success": True,
            "content": content,
            "file_path": file_path,
            "template_used": template.get("name", "默认模板"),
            "variables_used": variables,
        }

    def export_to_pdf(
        self,
        content: str,
        output_path: str,
        title: Optional[str] = None,
        author: Optional[str] = None
    ) -> bool:
        """导出为PDF文件

        Args:
            content: 文书内容
            output_path: 输出文件路径
            title: 文档标题（可选）
            author: 作者（可选）

        Returns:
            是否成功
        """
        if not REPORTLAB_AVAILABLE:
            print("错误: reportlab 未安装，无法生成PDF")
            return False

        try:
            # 创建PDF文档
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )

            # 获取样式
            styles = getSampleStyleSheet()

            # 创建自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='SimSun',
                fontSize=16,
                spaceAfter=30,
                alignment=1  # 居中
            )

            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='SimSun',
                fontSize=12,
                spaceAfter=12,
                leading=18
            )

            # 构建文档内容
            story = []

            # 添加标题
            if title:
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 0.25*inch))

            # 添加内容
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.replace('\n', '<br/>'), normal_style))

            # 生成PDF
            doc.build(story)
            return True

        except Exception as e:
            print(f"PDF生成错误: {e}")
            return False

    def export_to_docx(
        self,
        content: str,
        output_path: str,
        title: Optional[str] = None,
        author: Optional[str] = None
    ) -> bool:
        """导出为Word文档

        Args:
            content: 文书内容
            output_path: 输出文件路径
            title: 文档标题（可选）
            author: 作者（可选）

        Returns:
            是否成功
        """
        if not DOCX_AVAILABLE:
            print("错误: python-docx 未安装，无法生成Word文档")
            return False

        try:
            # 创建文档
            doc = Document()

            # 设置文档属性
            if title:
                doc.core_properties.title = title
            if author:
                doc.core_properties.author = author

            # 添加标题
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加内容
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para)
                    p.style.font.name = '宋体'
                    p.style.font.size = Pt(12)

            # 保存文档
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"Word文档生成错误: {e}")
            return False

    # ==================== 私有方法 ====================

    def _load_default_templates(self) -> Dict[str, Dict[str, Any]]:
        """加载默认模板"""
        return {
            "cover_letter_general": {
                "name": "通用套磁信模板",
                "category": "通用",
                "content_template": """尊敬的{professor_name}{professor_title}：

您好！

我是{student_name}，来自{student_university}{student_major}专业的一名学生。我对您的研究方向{professor_research_fields}非常感兴趣，特别是{research_match}方面的工作。

在本科期间，我学习了{relevant_courses}等课程，并在{research_experience}方面有了一定的研究经验。我的GPA为{student_gpa}，专业排名{student_ranking}。

通过阅读您的论文{mentioned_paper}，我对{specifiic_interest}有了更深入的理解。我相信我的背景和能力能够为您的课题组做出贡献。

希望能够有机会加入您的课题组继续深造，期待您的回复！

此致
敬礼！

{student_name}
{date}""",
                "variables": {
                    "professor_name": "导师姓名",
                    "professor_title": "职称（如：教授）",
                    "student_name": "学生姓名",
                    "student_university": "学生学校",
                    "student_major": "学生专业",
                    "professor_research_fields": "导师研究方向",
                    "research_match": "研究方向匹配点",
                    "relevant_courses": "相关课程",
                    "research_experience": "科研经历",
                    "student_gpa": "GPA",
                    "student_ranking": "专业排名",
                    "mentioned_paper": "提到的论文",
                    "specifiic_interest": "具体兴趣点",
                    "date": "日期"
                }
            },
            "personal_statement_general": {
                "name": "通用个人陈述模板",
                "category": "通用",
                "content_template": """个人陈述

一、学术背景
我来自{student_university}{student_major}专业，在校期间系统学习了{relevant_courses}等核心课程，取得了{student_gpa}的GPA，专业排名{student_ranking}。

二、科研经历
在科研方面，我主要参与了以下项目：
{research_experience_list}

通过这些项目，我掌握了{technical_skills}等技能，并培养了严谨的科研态度。

三、竞赛获奖
我曾参加{competition_list}等竞赛，获得了{competition_awards}。

四、研究兴趣与目标
我对{research_interests}方向有浓厚的兴趣，特别是{specifiic_interests}。我希望在研究生阶段能够深入研究{research_goals}。

五、职业规划
未来，我计划{career_plan}。

感谢您考虑我的申请！

{student_name}
{date}""",
                "variables": {
                    "student_university": "学生学校",
                    "student_major": "学生专业",
                    "relevant_courses": "相关课程",
                    "student_gpa": "GPA",
                    "student_ranking": "专业排名",
                    "research_experience_list": "科研经历列表",
                    "technical_skills": "技术技能",
                    "competition_list": "竞赛列表",
                    "competition_awards": "竞赛奖项",
                    "research_interests": "研究兴趣",
                    "specifiic_interests": "具体兴趣",
                    "research_goals": "研究目标",
                    "career_plan": "职业规划",
                    "student_name": "学生姓名",
                    "date": "日期"
                }
            },
            "recommendation_letter_general": {
                "name": "通用推荐信模板",
                "category": "通用",
                "content_template": """推荐信

致相关院系：

我很荣幸推荐{student_name}同学申请贵校的研究生项目。我是{recommender_name}，{recommender_title}，与{student_name}同学是{relationship}关系。

我认识{student_name}同学{known_period}。在此期间，他/她表现出了{outstanding_qualities}等优秀品质。

在学术方面，{student_name}同学学习态度认真，成绩优异（GPA: {student_gpa}，专业排名: {student_ranking}）。他/她掌握了{technical_skills}等专业技能。

在科研方面，{student_name}同学参与了{research_project}项目，表现出{research_abilities}等科研能力。

{student_name}同学具有{personal_qualities}等个人品质，是一位{overall_evaluation}的学生。

我强烈推荐{student_name}同学申请贵校的研究生项目，相信他/她能够在贵校取得优异的成绩。

此致
敬礼！

推荐人：{recommender_name}
{recommender_title}
{recommender_institution}
{date}""",
                "variables": {
                    "student_name": "学生姓名",
                    "recommender_name": "推荐人姓名",
                    "recommender_title": "推荐人职称",
                    "relationship": "关系",
                    "known_period": "认识时间",
                    "outstanding_qualities": "优秀品质",
                    "student_gpa": "GPA",
                    "student_ranking": "专业排名",
                    "technical_skills": "技术技能",
                    "research_project": "科研项目",
                    "research_abilities": "科研能力",
                    "personal_qualities": "个人品质",
                    "overall_evaluation": "总体评价",
                    "recommender_institution": "推荐人单位",
                    "date": "日期"
                }
            }
        }

    def _register_chinese_fonts(self):
        """注册中文字体"""
        try:
            # 尝试注册常用中文字体
            font_paths = [
                "C:/Windows/Fonts/simsun.ttc",  # Windows 宋体
                "/System/Library/Fonts/PingFang.ttc",  # macOS 苹方
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",  # Linux 文泉驿
            ]

            for font_path in font_paths:
                if Path(font_path).exists():
                    try:
                        pdfmetrics.registerFont(TTFont('SimSun', font_path))
                        print(f"中文字体已注册: {font_path}")
                        break
                    except Exception as e:
                        print(f"字体注册失败 {font_path}: {e}")
        except Exception as e:
            print(f"字体注册错误: {e}")

    def _select_template(
        self,
        student: Dict[str, Any],
        professor: Optional[Dict[str, Any]] = None,
        document_type: str = "套磁信",
        template_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """选择模板"""
        # TODO: 从数据库加载模板（如果template_id提供）
        # 目前使用默认模板

        template_key = f"{document_type.lower().replace(' ', '_')}_general"
        template = self.default_templates.get(template_key)

        if not template:
            # 回退到第一个可用模板
            template = next(iter(self.default_templates.values()))

        return template

    def _prepare_cover_letter_variables(
        self,
        student: Dict[str, Any],
        professor: Dict[str, Any],
        match_info: Optional[Dict[str, Any]] = None,
        custom_variables: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """准备套磁信变量"""
        variables = custom_variables.copy() if custom_variables else {}

        # 学生信息
        variables.update({
            "student_name": student.get("name", ""),
            "student_university": student.get("university", ""),
            "student_major": student.get("major", "相关专业"),
            "student_gpa": f"{student.get('gpa', '优秀')}",
            "student_ranking": f"前{student.get('gpa_ranking', '优秀')}%" if student.get("gpa_ranking") else "优秀",
        })

        # 导师信息
        variables.update({
            "professor_name": professor.get("name", ""),
            "professor_title": professor.get("title", "教授"),
            "professor_research_fields": ", ".join(professor.get("research_fields", [])),
        })

        # 匹配信息
        if match_info:
            research_match = match_info.get("details", {}).get("common_interests", [])
            if research_match:
                variables["research_match"] = ", ".join(research_match[:3])

        # 默认值
        defaults = {
            "relevant_courses": "材料力学、结构力学、复合材料力学等",
            "research_experience": "复合材料力学性能研究",
            "mentioned_paper": "您的近期工作",
            "specifiic_interest": "相关研究方向",
            "date": datetime.now().strftime("%Y年%m月%d日"),
        }

        for key, value in defaults.items():
            if key not in variables or not variables[key]:
                variables[key] = value

        return variables

    def _prepare_personal_statement_variables(
        self,
        student: Dict[str, Any],
        target_university: Optional[str] = None,
        target_major: Optional[str] = None,
        custom_variables: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """准备个人陈述变量"""
        variables = custom_variables.copy() if custom_variables else {}

        # 学生信息
        variables.update({
            "student_name": student.get("name", ""),
            "student_university": student.get("university", ""),
            "student_major": student.get("major", "相关专业"),
            "student_gpa": f"{student.get('gpa', '优秀')}",
            "student_ranking": f"前{student.get('gpa_ranking', '优秀')}%" if student.get("gpa_ranking") else "优秀",
        })

        # 科研经历
        research_experience = student.get("research_experience", [])
        if research_experience:
            exp_list = []
            for i, exp in enumerate(research_experience[:3], 1):
                if isinstance(exp, dict):
                    project = exp.get("project", "科研项目")
                    role = exp.get("role", "参与者")
                    duration = exp.get("duration", "")
                    exp_list.append(f"{i}. {project}（{role}，{duration}）")
            variables["research_experience_list"] = "\n".join(exp_list)
        else:
            variables["research_experience_list"] = "暂无系统科研经历"

        # 技能
        skills = student.get("skills", [])
        if skills:
            variables["technical_skills"] = ", ".join(skills[:5])
        else:
            variables["technical_skills"] = "相关专业技能"

        # 竞赛获奖
        competitions = student.get("competition_awards", [])
        if competitions:
            comp_list = []
            award_list = []
            for comp in competitions[:3]:
                if isinstance(comp, dict):
                    name = comp.get("name", "竞赛")
                    award = comp.get("award", "奖项")
                    comp_list.append(name)
                    award_list.append(award)
            variables["competition_list"] = "、".join(comp_list)
            variables["competition_awards"] = "、".join(award_list)
        else:
            variables["competition_list"] = "相关学科竞赛"
            variables["competition_awards"] = "优异成绩"

        # 目标院校和专业
        if target_university:
            variables["target_university"] = target_university
        if target_major:
            variables["target_major"] = target_major

        # 研究兴趣（从技能或研究方向中提取）
        research_interests = student.get("research_interests", []) or student.get("skills", [])
        if research_interests:
            variables["research_interests"] = ", ".join(research_interests[:3])

        # 默认值
        defaults = {
            "relevant_courses": "材料力学、结构力学、复合材料力学、有限元分析等",
            "specifiic_interests": "复合材料力学性能与智能材料结构",
            "research_goals": "智能材料与结构的前沿研究",
            "career_plan": "在学术界或工业界从事相关研究工作",
            "date": datetime.now().strftime("%Y年%m月%d日"),
        }

        for key, value in defaults.items():
            if key not in variables or not variables[key]:
                variables[key] = value

        return variables

    def _prepare_recommendation_letter_variables(
        self,
        student: Dict[str, Any],
        recommender: Dict[str, Any],
        relationship: str = "导师",
        custom_variables: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """准备推荐信变量"""
        variables = custom_variables.copy() if custom_variables else {}

        # 学生信息
        variables.update({
            "student_name": student.get("name", ""),
            "student_gpa": f"{student.get('gpa', '优秀')}",
            "student_ranking": f"前{student.get('gpa_ranking', '优秀')}%" if student.get("gpa_ranking") else "优秀",
        })

        # 推荐人信息
        variables.update({
            "recommender_name": recommender.get("name", "推荐人"),
            "recommender_title": recommender.get("title", "教授"),
            "recommender_institution": recommender.get("institution", student.get("university", "")),
            "relationship": relationship,
        })

        # 技能
        skills = student.get("skills", [])
        if skills:
            variables["technical_skills"] = ", ".join(skills[:5])

        # 科研经历
        research_experience = student.get("research_experience", [])
        if research_experience:
            # 取第一个项目作为示例
            first_exp = research_experience[0] if research_experience else {}
            if isinstance(first_exp, dict):
                variables["research_project"] = first_exp.get("project", "科研项目")

        # 默认值
        defaults = {
            "known_period": "两年",
            "outstanding_qualities": "勤奋好学、思维敏捷、责任心强",
            "technical_skills": "有限元分析、MATLAB编程、实验设计等",
            "research_abilities": "发现问题、分析问题和解决问题的能力",
            "personal_qualities": "团队合作精神、创新意识、坚韧不拔",
            "overall_evaluation": "综合素质优秀、具有很大发展潜力",
            "date": datetime.now().strftime("%Y年%m月%d日"),
        }

        for key, value in defaults.items():
            if key not in variables or not variables[key]:
                variables[key] = value

        return variables

    def _render_template(
        self,
        template: Dict[str, Any],
        variables: Dict[str, Any]
    ) -> str:
        """渲染模板"""
        content = template.get("content_template", "")

        # 替换变量
        for key, value in variables.items():
            placeholder = f"{{{key}}}"
            content = content.replace(placeholder, str(value))

        return content

    def _generate_file_path(
        self,
        student_id: Optional[int] = None,
        professor_id: Optional[int] = None,
        document_type: str = "文书"
    ) -> str:
        """生成文件路径"""
        # 生成唯一文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        random_str = str(uuid.uuid4())[:8]

        parts = []
        if student_id:
            parts.append(f"student_{student_id}")
        if professor_id:
            parts.append(f"professor_{professor_id}")
        parts.append(document_type.replace(" ", "_"))
        parts.append(timestamp)
        parts.append(random_str)

        filename = "_".join(parts) + ".txt"
        return f"documents/{filename}"